from PIL import Image, ImageDraw, ImageFont, ImageTk
from termcolor import colored
from colorama import init as colorama_init
from tkinter import messagebox
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import time
import datetime
import argparse
import tkinter as tk

parser = argparse.ArgumentParser()

parser.add_argument("-c", "--console", action="store_true", help="Start program in console mode")
parser.add_argument("-t", "--test", action="store_true", help="Generate test receipt")

args = parser.parse_args()

colorama_init()

# Config
CERTIFICATE_TYPE = ('png') # ('pdf', 'png')
TEMPLATE_PATH = os.path.join('template')
NAME_LIST_PATH = os.path.join('list')
NAME_LIST_DICT = {}
NAME_FORMAT = 'capitalize' # capitalize, upper, lower
PNG_PATH = 'png'
PDF_PATH = 'pdf'
FONT_PATH = os.path.join('font')
FONT_TYPE = []
FONT_SIZE = 18 #20
FONT_SIZE_ADJUSTMENT = -2
FONT_COLOR = '#0000'
WIDTH = 379 #876
HEIGHT = 237 #620
HORIZONTAL_OFFSET = 170
HORIZONTAL_OFFSET_ADJUSTMENT = 0
VERTICAL_OFFSET = 100
VERTICAL_OFFSET_ADJUSTMENT = -5
ERROR_COLOR = 'red'
SUCCESS_COLOR = 'green'

def error(args):
    '''
    Prints out args in red color
    '''
    global ERROR_COLOR
    print(colored(args, ERROR_COLOR))

def success(args):
    '''
    Prints out args in green color
    '''
    global SUCCESS_COLOR
    print(colored(args, SUCCESS_COLOR))

def format_time(solve_time):
    hours = int(solve_time // 3600)
    solve_time %= 3600
    minutes = int(solve_time // 60)
    seconds = solve_time % 60
    return '{:02d}:{:02d}:{:05.2f}'.format(hours, minutes, seconds)

def create_certificate(name, phone_number, address1, address2=None, address3=None):
    '''
    Generate certificate for name
    '''
    global FONT_TYPE, FONT_SIZE, FONT_COLOR, HEIGHT, WIDTH, PNG_PATH, PDF_PATH, CERTIFICATE_TYPE, FONT_SIZE_ADJUSTMENT, HORIZONTAL_OFFSET, HORIZONTAL_OFFSET_ADJUSTMENT, VERTICAL_OFFSET, VERTICAL_OFFSET_ADJUSTMENT
    print(f"[!] Starting receipt creation for {name}")

    # Template Certificate name

    im = Image.open(os.path.join(TEMPLATE_PATH, 'template' + '.png')).convert('RGB')

    # Choosing Color for each certificate
    text_color = FONT_COLOR
    # Inserting the name to the certificate
    d = ImageDraw.Draw(im)
    font1 = ImageFont.truetype(FONT_TYPE[0], FONT_SIZE)
    font2 = ImageFont.truetype(FONT_TYPE[1], FONT_SIZE)
    t1 = f'{name} ({phone_number})'
    if not address2 or not address3:
        t2 = address1
    else:
        t2 = f'\n{address1}\n{address2}\n{address3}'
    w1, h1 = d.textsize(t1, font=font1)
    w2, h2 = d.textsize(t2, font=font2)
    # print(f'{w=} {h=}')
    # location1 = (424.5 - w1, 172)
    # location2 = (424.5 - w2, 172)
    location1 = (WIDTH - 15 - w1, 140)
    location2 = (WIDTH - 15 - w2, 140+25)
    font = ImageFont.truetype(FONT_TYPE[0], font1.size + FONT_SIZE_ADJUSTMENT)
    while (location1[0] < 0 or location1[0] > WIDTH):
        font = ImageFont.truetype(FONT_TYPE[0], font.size + FONT_SIZE_ADJUSTMENT)
        w, h = d.textsize(name, font=font)
        location1 = ((WIDTH-w) / 2 + HORIZONTAL_OFFSET_ADJUSTMENT, (HEIGHT-h) / 2 + VERTICAL_OFFSET_ADJUSTMENT)
    while (location2[0] < 0 or location2[0] > WIDTH):
        font = ImageFont.truetype(FONT_TYPE[1], font.size + FONT_SIZE_ADJUSTMENT)
        w, h = d.textsize(name, font=font)
        location2 = ((WIDTH-w) / 2 + HORIZONTAL_OFFSET_ADJUSTMENT, (HEIGHT-h) / 2 + VERTICAL_OFFSET_ADJUSTMENT)
    print(location1)
    print(location2)
    d.text(location1, t1, fill=text_color, font=font1, align='right')
    d.text(location2, t2, fill=text_color, font=font2, align='right')

    # Save the certificate name

    filename = f'{name}-{format_datetime(datetime.datetime.now())}'

    if 'png' in CERTIFICATE_TYPE:
        im.save(os.path.join(PNG_PATH, f'{filename}.png'))
    if 'pdf' in CERTIFICATE_TYPE:
        im.save(os.path.join(PDF_PATH, f'{filename}.pdf'))
        # im.save(os.path.join(PDF_PATH, os.path.join('{}-{}.pdf'.format(name, format_datetime(datetime.datetime.now())))))
    success(f"[o] Successfully created receipt for {name}")
    return filename

def save_to_docx(filename):
    '''
    Save png file into a docx format to print
    '''
    global PNG_PATH
    document = Document()
    document.add_picture(os.path.join(PNG_PATH, f'{filename}.png'), width=Inches(3.95), height=Inches(2.47))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for section in document.sections:
        section.top_margin = Inches(0.12)
    document.save(os.path.join(PNG_PATH, 'Print Out.docx'))
    success(f"[o] Successfully created receipt document for {filename}")

def init():
    '''
    Checks required paths, and create them if doesn't exists
    '''
    global TEMPLATE_PATH, PNG_PATH, PDF_PATH, NAME_LIST_PATH, CERTIFICATE_TYPE

    if not os.path.exists(NAME_LIST_PATH):
        error('Name list path is not found!')
        os.mkdir(NAME_LIST_PATH)
        success(f'Created template path in {NAME_LIST_PATH}')
    else:
        success(f'Template path is found on {NAME_LIST_PATH}')

    if not os.path.exists(PNG_PATH) and 'png' in CERTIFICATE_TYPE:
        error('PNG path is not found!')
        os.mkdir(PNG_PATH)
        success(f'Created template path in {PNG_PATH}')
    else:
        success(f'Template path is found on {PNG_PATH}')

    if not os.path.exists(PDF_PATH) and 'pdf' in CERTIFICATE_TYPE:
        error('PDF path is not found!')
        os.mkdir(PDF_PATH)
        success(f'Created PDF path in {PDF_PATH}')
    else:
        success(f'PDF path is found on {PDF_PATH}')

    if not os.path.exists(TEMPLATE_PATH):
        error('Template path is not found!')
        os.mkdir(TEMPLATE_PATH)
        success(f'Created template path in {TEMPLATE_PATH}')
    else:
        success(f'Template path is found on {TEMPLATE_PATH}')

    if not os.path.exists(FONT_PATH):
        error('Font path is not found!')
        os.mkdir(FONT_PATH)
        success(f'Created font path in {FONT_PATH}')
    else:
        success(f'Font path is found on {FONT_PATH}')

def load_name_list():
    '''
    Loads name list file
    '''
    global NAME_LIST_PATH, NAME_LIST_DICT
    file_list = os.listdir(NAME_LIST_PATH)
    for f in file_list:
        with open(os.path.join(NAME_LIST_PATH, f)) as file:
            NAME_LIST_DICT[f.replace('.csv', '')] = file.read().splitlines()

def load_font_type():
    '''
    Loads font type from font path
    '''
    global FONT_PATH, FONT_TYPE
    try:
        fonts = list(filter(lambda x: x.endswith('.ttf') or x.endswith('.otf'), os.listdir(FONT_PATH)))
        print(fonts)
        for font in fonts:
            FONT_TYPE.append(os.path.join(FONT_PATH, font))
            print(FONT_TYPE)
        success(f'{FONT_TYPE} loaded from {FONT_PATH}')
    except:
        error(f'Cannot find font file in {FONT_PATH}')

def format_name_upper(name):
    '''
    Returns formatted name as FIRSTNAME LASTNAME
    '''
    return name.upper().replace('"', '')

def format_name_lower(name):
    '''
    Returns formatted name as firstname lastname
    '''
    return name.lower().replace('"', '')

def format_name_capitalize(name):
    '''
    Returns formatted name as Fisrtname Lastname
    '''
    return ' '.join(list(map(lambda x: x.capitalize(), name.split(' ')))).replace('"', '')

def format_datetime(datetime):
    return ''.join([ '.' if x == ':' else x for x in list(str(datetime))])

def format_phone_number(phone_number):
    phone_number_list = []

    for i in range(len(phone_number)):
        if i != 0 and i % 4 == 0:
            phone_number_list.append('-')
        phone_number_list.append(phone_number[i])
    return ''.join(phone_number_list)

def get_input():
    name = input('Input recipient name: ')
    phone_number = format_phone_number(input('Input recipient phone number: '))
    address1 = input('Input recipient address 1: ')
    address2 = input('Input recipient address 2: ')
    address3 = input('Input recipient address 3: ')
    return name, phone_number, address1, address2, address3

def on_click():
    name = entry_name.get()
    phone_number = format_phone_number(entry_phone_number.get())
    address = text_address.get("1.0", tk.END)

    filename = create_certificate(name=name, phone_number=phone_number, address1=address)
    save_to_docx(filename)
    messagebox.showinfo("Success","Succesfully generated receipt!")

if __name__ == '__main__':
    init()
    load_font_type()
    # load_name_list()
    if not args.console:
        root = tk.Tk()
        root.geometry("500x200")
        root.wm_iconbitmap('logo.ico')
        root.wm_title('Theresia Crochet Receipt Generator v2')

        frame_name = tk.Frame(root)
        label_name = tk.Label(frame_name, text='Input recipient name: ')
        entry_name = tk.Entry(frame_name, width=49)
        frame_name.pack(side= tk.TOP, fill= tk.X, padx = 5, pady= 5)
        label_name.pack(side= tk.LEFT)
        entry_name.pack(side= tk.RIGHT, fill = tk.X)

        frame_phone_number = tk.Frame(root)
        label_phone_number = tk.Label(frame_phone_number, text='Input recipient phone number: ')
        entry_phone_number = tk.Entry(frame_phone_number, width=49)
        frame_phone_number.pack(side= tk.TOP, fill= tk.X, padx = 5, pady= 5)
        label_phone_number.pack(side= tk.LEFT)
        entry_phone_number.pack(side= tk.RIGHT, fill = tk.X)

        frame_address = tk.Frame(root)
        label_address = tk.Label(frame_address, text='Input recipient address: ')
        text_address = tk.Text(frame_address, height=3, width=37)
        frame_address.pack(side= tk.TOP, fill= tk.X, padx = 5, pady= 5)
        label_address.pack(side= tk.LEFT)
        text_address.pack(side= tk.RIGHT, fill = tk.X)

        frame_button = tk.Frame(root)
        button_submit = tk.Button(frame_button, text='Submit', command=on_click, bg="#20bebe", fg="white", height=2, width=15)
        frame_button.pack(side= tk.TOP, fill= tk.X, padx = 5, pady= 20)
        button_submit.pack(side= tk.TOP)

        root.mainloop()
    else:
        if args.test:
            name = 'Dummy Data'
            phone_number = '1234 5678 9101'
            address1 = 'Jl. Jalan Timur no. 14'
            address2 = 'Kecamatan - Kelurahan'
            address3 = 'Kota - 11111'
        else:
            name, phone_number, address1, address2, address3 = get_input()
            
        counter = 1
        print("[+] Starting receipt generation...")
        start = time.time()
        if NAME_FORMAT == 'upper':
            f_name = format_name_upper(name)
        elif NAME_FORMAT == 'lower':
            f_name = format_name_lower(name)
        else :
            f_name = format_name_capitalize(name)
        create_certificate(f_name, phone_number, address1, address2, address3)
        counter += 1
        # for key in NAME_LIST_DICT:
        #     if 'png' in CERTIFICATE_TYPE:
        #         try:
        #             os.mkdir(os.path.join(PNG_PATH, key))
        #         except:
        #             pass
        #     if 'pdf' in CERTIFICATE_TYPE:
        #         try:
        #             os.mkdir(os.path.join(PDF_PATH, key))
        #         except:
        #             pass
        #     names = NAME_LIST_DICT[key]
        #     for name in names:
        #         if NAME_FORMAT == 'upper':
        #             f_name = format_name_upper(name)
        #         elif NAME_FORMAT == 'lower':
        #             f_name = format_name_lower(name)
        #         else :
        #             f_name = format_name_capitalize(name)
        #         create_certificate(key, f_name)
        #         counter += 1
        end = time.time()
        duration = end - start
        print()
        print(f'The process took {format_time(duration)} to complete\n')
